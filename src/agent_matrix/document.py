"""文档解析模块 - 支持 Word(.docx) 和 PDF 文档解析"""

from __future__ import annotations

from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional


class DocumentParser(ABC):
    """文档解析器基类"""

    @property
    @abstractmethod
    def supported_extensions(self) -> List[str]:
        """支持的扩展名"""
        ...

    @abstractmethod
    def parse(self, file_path: str) -> str:
        """解析文档并返回文本内容

        Args:
            file_path: 文档路径

        Returns:
            文档文本内容
        """
        ...

    def can_parse(self, file_path: str) -> bool:
        """检查是否能够解析此文件

        Args:
            file_path: 文件路径

        Returns:
            是否可以解析
        """
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions


class PDFParser(DocumentParser):
    """PDF 文档解析器"""

    @property
    def supported_extensions(self) -> List[str]:
        return [".pdf"]

    def parse(self, file_path: str) -> str:
        """解析 PDF 文档

        Args:
            file_path: PDF 文件路径

        Returns:
            文档文本内容
        """
        from pypdf import PdfReader

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        reader = PdfReader(str(path))
        text_parts = []

        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text:
                text_parts.append(f"--- 第 {page_num} 页 ---\n{text}")

        return "\n\n".join(text_parts)


class WordParser(DocumentParser):
    """Word 文档解析器"""

    @property
    def supported_extensions(self) -> List[str]:
        return [".docx"]

    def parse(self, file_path: str) -> str:
        """解析 Word 文档

        Args:
            file_path: Word 文件路径

        Returns:
            文档文本内容
        """
        from docx import Document

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        doc = Document(str(path))
        text_parts = []

        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 提取表格
        for table_idx, table in enumerate(doc.tables, 1):
            table_text = [f"--- 表格 {table_idx} ---"]
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_text.append(" | ".join(cells))
            if len(table_text) > 1:
                text_parts.append("\n".join(table_text))

        return "\n\n".join(text_parts)


class DocumentService:
    """文档服务统一入口"""

    def __init__(self):
        self._parsers: List[DocumentParser] = [
            PDFParser(),
            WordParser(),
        ]

    def parse(self, file_path: str) -> str:
        """解析文档

        Args:
            file_path: 文档路径

        Returns:
            文档文本内容

        Raises:
            ValueError: 不支持的文档格式
            FileNotFoundError: 文件不存在
        """
        ext = Path(file_path).suffix.lower()

        for parser in self._parsers:
            if parser.can_parse(file_path):
                return parser.parse(file_path)

        raise ValueError(f"不支持的文档格式: {ext}，支持的格式: {self.supported_formats()}")

    def can_parse(self, file_path: str) -> bool:
        """检查是否能够解析此文件"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions

    @property
    def supported_extensions(self) -> List[str]:
        """获取所有支持的扩展名"""
        extensions = []
        for parser in self._parsers:
            extensions.extend(parser.supported_extensions)
        return list(set(extensions))

    def supported_formats(self) -> str:
        """获取支持格式的描述"""
        return ", ".join(f"*{ext}" for ext in self.supported_extensions)


# 全局服务实例
_document_service: Optional[DocumentService] = None


def get_document_service() -> DocumentService:
    """获取文档服务单例"""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service


def parse_document(file_path: str) -> str:
    """快捷函数：解析单个文档"""
    return get_document_service().parse(file_path)


def parse_documents(file_paths: List[str]) -> List[tuple[str, str]]:
    """解析多个文档

    Args:
        file_paths: 文档路径列表

    Returns:
        [(file_path, content), ...] 元组列表
    """
    service = get_document_service()
    results = []
    for path in file_paths:
        try:
            content = service.parse(path)
            results.append((path, content))
        except Exception as e:
            results.append((path, f"解析失败: {e}"))
    return results
